"""Document template handler - generates documents from templates"""
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Optional
import os

from ..models.models import Employee, Company, DocumentConfig


class TemplateHandler:
    """Handles document generation from templates"""

    # French month names for date formatting
    MONTHS = {
        1: "Janvier", 2: "Février", 3: "Mars", 4: "Avril", 5: "Mai", 6: "Juin",
        7: "Juillet", 8: "Août", 9: "Septembre", 10: "Octobre", 11: "Novembre", 12: "Décembre"
    }

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize template handler

        Args:
            template_path: Path to custom template directory
        """
        self.template_dir = Path(template_path) if template_path else Path("templates")
        self.template_dir.mkdir(parents=True, exist_ok=True)

    def _resolve_output_path(self, output_path: str) -> str:
        """
        Resolve output path to absolute path
        
        Args:
            output_path: Relative or absolute output path
            
        Returns:
            Absolute path as string
        """
        path = Path(output_path)
        if not path.is_absolute():
            # If relative, resolve from current working directory
            path = Path.cwd() / path
        return str(path)

    def _add_header_footer_images(self, doc: Document):
        """
        Add logo to header (top left) and footer info to footer (bottom left)
        
        Args:
            doc: Document object to modify
        """
        try:
            sections = doc.sections
            for section in sections:
                # Add logo to header (top left)
                header = section.header
                header_table = header.add_table(rows=1, cols=1, width=Inches(6))
                header_table.autofit = False
                header_cell = header_table.rows[0].cells[0]
                
                logo_path = Path("templates/logo.png")
                if logo_path.exists():
                    header_paragraph = header_cell.paragraphs[0]
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = header_paragraph.add_run()
                    run.add_picture(str(logo_path), width=Inches(1.2))
                
                # Add footer info to footer (bottom left)
                footer = section.footer
                footer_table = footer.add_table(rows=1, cols=1, width=Inches(6))
                footer_table.autofit = False
                footer_cell = footer_table.rows[0].cells[0]
                
                footer_info_path = Path("templates/Footer.png")
                if footer_info_path.exists():
                    footer_paragraph = footer_cell.paragraphs[0]
                    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = footer_paragraph.add_run()
                    run.add_picture(str(footer_info_path), width=Inches(2.5))
        except Exception as e:
            print(f"Warning: Could not add header/footer images: {e}")

    def format_french_date(self, date: datetime) -> str:
        """Format date in French: '23 Mars 2026'"""
        return f"{date.day} {self.MONTHS[date.month]} {date.year}"

    def get_gender_title(self, gender: str, position: str = "employé") -> tuple:
        """
        Get proper French gender title and adjective

        Args:
            gender: "Madame" or "Monsieur"
            position: Type of position/action

        Returns:
            Tuple of (title, adjective_masculine, adjective_feminine)
        """
        if gender.lower() == "madame":
            return ("Mme", False, True)
        else:
            return ("M.", True, False)

    def create_placeholders_dict(self, config: DocumentConfig) -> Dict[str, str]:
        """
        Create dictionary of placeholders to replace

        Args:
            config: DocumentConfig with all document data

        Returns:
            Dictionary mapping placeholder names to values
        """
        gender_title, is_male, is_female = self.get_gender_title(config.employee.gender)

        return {
            "{{company_name}}": config.company.name,
            "{{legal_id}}": config.company.legal_id,
            "{{representative_name}}": config.company.representative_name,
            "{{company_address}}": config.company.address,
            "{{employee_name}}": config.employee.name,
            "{{employee_gender_title}}": gender_title,
            "{{cin}}": config.employee.cin,
            "{{cin_place}}": config.employee.cin_place,
            "{{cin_date}}": self.format_french_date(config.employee.cin_date),
            "{{position}}": config.employee.position,
            "{{start_date}}": self.format_french_date(config.employee.start_date),
            "{{current_date}}": self.format_french_date(config.current_date),
            "{{reference}}": config.reference,
            "{{city}}": config.company.city,
            # Gender-specific variants
            "{{employe_e}}": "employée" if is_female else "employé",
            "{{employe_male}}": "employé",
            "{{employe_female}}": "employée",
        }

    def replace_text_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        """Replace placeholders in a paragraph"""
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # Replace in full paragraph text
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                # If placeholder spans multiple runs, reconstruct
                full_text = "".join([run.text for run in paragraph.runs])
                if placeholder in full_text:
                    # Clear runs and rebuild
                    for run in paragraph.runs:
                        run.text = ""
                    new_text = full_text.replace(placeholder, value)
                    paragraph.text = new_text

    def replace_text_in_table(self, table, replacements: Dict[str, str]):
        """Replace placeholders in table cells"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_text_in_paragraph(paragraph, replacements)

    def create_attestation_document(self, config: DocumentConfig, output_path: str) -> str:
        """
        Create Attestation de Travail document

        Args:
            config: Document configuration
            output_path: Path where to save the document

        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Add header and footer with images
        self._add_header_footer_images(doc)

        placeholders = self.create_placeholders_dict(config)
        gender_title, is_male, is_female = self.get_gender_title(config.employee.gender)

        # Header: City and Date (top right) - Roboto 12 bold
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(f"{config.company.city}, le {placeholders['{{current_date}}']}")
        header_run.font.name = 'Roboto'
        header_run.font.size = Pt(12)
        header_run.font.bold = True

        # Reference - Roboto 12 not bold
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"Réf / DG : {config.reference}")
        ref_run.font.name = 'Roboto'
        ref_run.font.size = Pt(12)
        ref_run.font.bold = False

        # Add spacing
        doc.add_paragraph()

        # Title - Roboto 16 bold
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("ATTESTATION DE TRAVAIL")
        title_run.font.name = 'Roboto'
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        # Add spacing
        doc.add_paragraph()

        # Body text - Roboto 14
        body_text = f"""La société {placeholders['{{company_name}}']} M.F. {placeholders['{{legal_id}}']}, représentée par sa fondée de pouvoir {placeholders['{{employee_gender_title}}']}, {placeholders['{{representative_name}}']}, sise au {placeholders['{{company_address}}']}, atteste par la présente {placeholders['{{employee_gender_title}}']} {placeholders['{{employee_name}}']}, titulaire de la carte d'identité nationale N° {placeholders['{{cin}}']} délivrée à {placeholders['{{cin_place}}']} le {placeholders['{{cin_date}}']}, est {placeholders['{{employe_e}}']} au sein de notre société en tant que {placeholders['{{position}}']} depuis le {placeholders['{{start_date}}']}."""

        body_para = doc.add_paragraph(body_text)
        body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in body_para.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(14)

        # Add spacing
        doc.add_paragraph()

        # Second body paragraph - Roboto 14
        closing_text = "Cette attestation est délivrée à l'intéressé sur sa demande pour servir et valoir ce que de droit."
        closing_para = doc.add_paragraph(closing_text)
        closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in closing_para.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(14)

        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Signature block (right aligned) - Roboto 14 bold
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_para.add_run(placeholders['{{representative_name}}'])
        sig_run.font.name = 'Roboto'
        sig_run.font.size = Pt(14)
        sig_run.font.bold = True
        
        sig_para2 = doc.add_paragraph()
        sig_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run2 = sig_para2.add_run("Directrice Générale")
        sig_run2.font.name = 'Roboto'
        sig_run2.font.size = Pt(14)
        sig_run2.font.bold = True
        
        sig_para3 = doc.add_paragraph()
        sig_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run3 = sig_para3.add_run("PLW Tunisia")
        sig_run3.font.name = 'Roboto'
        sig_run3.font.size = Pt(14)
        sig_run3.font.bold = True

        # Save document
        output_file = Path(self._resolve_output_path(output_path))
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Verify directory is writable
        if not os.access(output_file.parent, os.W_OK):
            raise Exception(f"No write permission for directory: {output_file.parent}")
        
        # Save the document
        doc.save(str(output_file))
        
        # Verify file was created
        if not output_file.exists():
            raise Exception(f"File was not created at: {output_file}")
        
        return str(output_file)

    def create_ordre_mission_document(self, config: DocumentConfig, output_path: str, country: str = "Allemagne", nationality: str = "", address: str = "", mission_end_date = None) -> str:
        """
        Create Ordre de Mission document with dynamic location based on country

        Args:
            config: Document configuration
            output_path: Path where to save the document
            country: Destination country (Allemagne or France)
            nationality: Employee nationality
            address: Employee address
            mission_end_date: End date of mission

        Returns:
            Path to generated document
        """
        doc = Document()
        
        # Add header and footer with images
        self._add_header_footer_images(doc)

        # Get gender-specific text
        gender_title, is_male, is_female = self.get_gender_title(config.employee.gender)
        
        # Format dates
        birth_date_str = self.format_french_date(config.employee.cin_date)
        mission_start_str = self.format_french_date(config.employee.start_date)
        mission_end_str = self.format_french_date(mission_end_date) if mission_end_date else mission_start_str

        # Header: City and Date (top right) - Roboto 12 bold
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run(f"{config.company.city}, le {self.format_french_date(config.current_date)}")
        header_run.font.name = 'Roboto'
        header_run.font.size = Pt(12)
        header_run.font.bold = True

        # Reference - Roboto 12 not bold
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"Réf / DG : {config.reference}")
        ref_run.font.name = 'Roboto'
        ref_run.font.size = Pt(12)
        ref_run.font.bold = False

        # Add spacing
        doc.add_paragraph()

        # Title - Roboto 16 bold
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("ORDRE DE MISSION")
        title_run.font.name = 'Roboto'
        title_run.font.size = Pt(16)
        title_run.font.bold = True

        # Add spacing
        doc.add_paragraph()

        # Main content - Roboto 14
        # Personal info section
        intro_para = doc.add_paragraph()
        intro_text = f"{gender_title} {config.employee.name} de nationalité {nationality}, né(e) le {birth_date_str} à {config.employee.cin_place},"
        intro_run = intro_para.add_run(intro_text)
        intro_run.font.name = 'Roboto'
        intro_run.font.size = Pt(14)

        # CIN and address
        cin_para = doc.add_paragraph()
        cin_text = f"titulaire de la Carte d'Identité Nationale numéro {config.employee.cin} délivrée à {config.employee.cin_place} et demeurant au {address}."
        cin_run = cin_para.add_run(cin_text)
        cin_run.font.name = 'Roboto'
        cin_run.font.size = Pt(14)

        # Position
        position_para = doc.add_paragraph()
        position_text = f"Fonction : {config.employee.position} depuis le {mission_start_str}."
        position_run = position_para.add_run(position_text)
        position_run.font.name = 'Roboto'
        position_run.font.size = Pt(14)

        # Location section - varies by country
        location_para = doc.add_paragraph()
        location_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if country.lower() == "allemagne":
            location_text = f"À : Planisware Deutschland GMBH, Leonrodstraße 52 - 58, 80636 München, Germany, une entreprise qui fait partie du groupe Planisware pour suivre une formation du {mission_start_str} au {mission_end_str}."
        else:  # France
            location_text = f"À : Planisware SA, 200 Av. de Paris, 92320 Châtillon, France une entreprise qui fait partie du groupe Planisware pour suivre une formation du {mission_start_str} au {mission_end_str}."

        location_run = location_para.add_run(location_text)
        location_run.font.name = 'Roboto'
        location_run.font.size = Pt(14)

        # Mission purpose
        if country.lower() == "allemagne":
            mission_purpose = "Pour la mission : Formation professionnel, renforcer la coopération entre les deux équipes (Tunis et Allemagne) et rencontre clients Planisware."
        else:  # France
            mission_purpose = "Pour la mission : Formation professionnel, renforcer la coopération entre les deux équipes (Tunis et France) et rencontre clients Planisware."

        purpose_para = doc.add_paragraph(mission_purpose)
        purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in purpose_para.runs:
            run.font.name = 'Roboto'
            run.font.size = Pt(14)

        # Add spacing before signature
        doc.add_paragraph()
        doc.add_paragraph()

        # Signature block - Roboto 14 bold
        sig_para = doc.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run = sig_para.add_run(config.company.representative_name if config.company.representative_name else "Mme Lobna MILED")
        sig_run.font.name = 'Roboto'
        sig_run.font.size = Pt(14)
        sig_run.font.bold = True
        
        sig_para2 = doc.add_paragraph()
        sig_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run2 = sig_para2.add_run("Directrice Générale")
        sig_run2.font.name = 'Roboto'
        sig_run2.font.size = Pt(14)
        sig_run2.font.bold = True
        
        sig_para3 = doc.add_paragraph()
        sig_para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig_run3 = sig_para3.add_run("PLW Tunisia")
        sig_run3.font.name = 'Roboto'
        sig_run3.font.size = Pt(14)
        sig_run3.font.bold = True

        # Save document
        output_file = Path(self._resolve_output_path(output_path))
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Verify directory is writable
        if not os.access(output_file.parent, os.W_OK):
            raise Exception(f"No write permission for directory: {output_file.parent}")
        
        # Save the document
        doc.save(str(output_file))
        
        # Verify file was created
        if not output_file.exists():
            raise Exception(f"File was not created at: {output_file}")
        
        return str(output_file)


def gender_title_full(title: str) -> str:
    """Convert short gender title to full"""
    return "Madame" if title == "Mme" else "Monsieur"
