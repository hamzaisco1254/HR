"""Document template handler - generates DOCX documents in memory."""
import io
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import Employee, Company, DocumentConfig

# Resolve the project root (one level above this api/ directory)
_API_DIR = os.path.dirname(os.path.abspath(__file__))
_ROOT_DIR = os.path.dirname(_API_DIR)
_TEMPLATES_DIR = os.path.join(_ROOT_DIR, 'templates')


class TemplateHandler:
    """Handles DOCX document generation."""

    MONTHS = {
        1: 'Janvier', 2: 'Février', 3: 'Mars', 4: 'Avril',
        5: 'Mai', 6: 'Juin', 7: 'Juillet', 8: 'Août',
        9: 'Septembre', 10: 'Octobre', 11: 'Novembre', 12: 'Décembre',
    }

    def __init__(self, template_path: Optional[str] = None):
        # Allow override; otherwise use the resolved project templates dir
        self.template_dir = template_path or _TEMPLATES_DIR

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _logo_path(self) -> Optional[str]:
        p = os.path.join(self.template_dir, 'logo.png')
        return p if os.path.exists(p) else None

    def _footer_path(self) -> Optional[str]:
        p = os.path.join(self.template_dir, 'Footer.png')
        return p if os.path.exists(p) else None

    def _add_header_footer_images(self, doc: Document):
        try:
            logo = self._logo_path()
            footer_img = self._footer_path()
            for section in doc.sections:
                if logo:
                    header = section.header
                    tbl = header.add_table(rows=1, cols=1, width=Inches(6))
                    tbl.autofit = False
                    para = tbl.rows[0].cells[0].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.add_run().add_picture(logo, width=Inches(1.2))

                if footer_img:
                    footer = section.footer
                    tbl = footer.add_table(rows=1, cols=1, width=Inches(6))
                    tbl.autofit = False
                    para = tbl.rows[0].cells[0].paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.add_run().add_picture(footer_img, width=Inches(2.5))
        except Exception as e:
            print(f"Warning: could not add header/footer images: {e}")

    def format_french_date(self, date: datetime) -> str:
        return f"{date.day} {self.MONTHS[date.month]} {date.year}"

    def get_gender_title(self, gender: str):
        if gender.lower() == 'madame':
            return ('Mme', False, True)
        return ('M.', True, False)

    def _run(self, para, text: str, bold: bool = False, size: int = 14):
        run = para.add_run(text)
        run.font.name = 'Roboto'
        run.font.size = Pt(size)
        run.font.bold = bold
        return run

    # ------------------------------------------------------------------
    # Public document generators — return BytesIO
    # ------------------------------------------------------------------

    def create_attestation_document(self, config: DocumentConfig, output_path: str = None) -> io.BytesIO:
        """Generate Attestation de Travail. Returns a BytesIO buffer."""
        doc = Document()
        self._add_header_footer_images(doc)

        gender_title, is_male, is_female = self.get_gender_title(config.employee.gender)
        employe_e = 'employée' if is_female else 'employé'
        current_date_str = self.format_french_date(config.current_date)

        # Date + city (top right)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(p, f"{config.company.city}, le {current_date_str}", bold=True, size=12)

        # Reference
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(p, f"Réf / DG : {config.reference}", bold=False, size=12)

        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p, 'ATTESTATION DE TRAVAIL', bold=True, size=16)

        doc.add_paragraph()

        # Body
        body = (
            f"La société {config.company.name} M.F. {config.company.legal_id}, "
            f"représentée par sa fondée de pouvoir {gender_title}, "
            f"{config.company.representative_name}, "
            f"sise au {config.company.address}, "
            f"atteste par la présente {gender_title} {config.employee.name}, "
            f"titulaire de la carte d'identité nationale N° {config.employee.cin} "
            f"délivrée à {config.employee.cin_place} le "
            f"{self.format_french_date(config.employee.cin_date)}, "
            f"est {employe_e} au sein de notre société en tant que "
            f"{config.employee.position} depuis le "
            f"{self.format_french_date(config.employee.start_date)}."
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, body, size=14)

        doc.add_paragraph()

        # Closing
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(
            p,
            "Cette attestation est délivrée à l'intéressé sur sa demande "
            "pour servir et valoir ce que de droit.",
            size=14,
        )

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        for line in [config.company.representative_name, 'Directrice Générale', 'PLW Tunisia']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._run(p, line, bold=True, size=14)

        # Save to output_path if given (backward compat), always return BytesIO
        buf = io.BytesIO()
        doc.save(buf)
        if output_path:
            doc.save(output_path)
        buf.seek(0)
        return buf

    def create_ordre_mission_document(
        self,
        config: DocumentConfig,
        output_path: str = None,
        country: str = 'Allemagne',
        nationality: str = '',
        address: str = '',
        mission_end_date=None,
    ) -> io.BytesIO:
        """Generate Ordre de Mission. Returns a BytesIO buffer."""
        doc = Document()
        self._add_header_footer_images(doc)

        gender_title, is_male, is_female = self.get_gender_title(config.employee.gender)
        birth_date_str = self.format_french_date(config.employee.cin_date)
        mission_start_str = self.format_french_date(config.employee.start_date)
        mission_end_str = (
            self.format_french_date(mission_end_date) if mission_end_date else mission_start_str
        )

        # Date + city
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(p, f"{config.company.city}, le {self.format_french_date(config.current_date)}", bold=True, size=12)

        # Reference
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(p, f"Réf / DG : {config.reference}", bold=False, size=12)

        doc.add_paragraph()

        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p, 'ORDRE DE MISSION', bold=True, size=16)

        doc.add_paragraph()

        # Personal info
        p = doc.add_paragraph()
        self._run(
            p,
            f"{gender_title} {config.employee.name} de nationalité {nationality}, "
            f"né(e) le {birth_date_str} à {config.employee.cin_place},",
            size=14,
        )

        p = doc.add_paragraph()
        self._run(
            p,
            f"titulaire de la Carte d'Identité Nationale numéro {config.employee.cin} "
            f"délivrée à {config.employee.cin_place} et demeurant au {address}.",
            size=14,
        )

        p = doc.add_paragraph()
        self._run(
            p,
            f"Fonction : {config.employee.position} depuis le {mission_start_str}.",
            size=14,
        )

        # Location
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if country.lower() == 'allemagne':
            loc = (
                f"À : Planisware Deutschland GMBH, Leonrodstraße 52 - 58, 80636 München, Germany, "
                f"une entreprise qui fait partie du groupe Planisware pour suivre une formation "
                f"du {mission_start_str} au {mission_end_str}."
            )
        else:
            loc = (
                f"À : Planisware SA, 200 Av. de Paris, 92320 Châtillon, France "
                f"une entreprise qui fait partie du groupe Planisware pour suivre une formation "
                f"du {mission_start_str} au {mission_end_str}."
            )
        self._run(p, loc, size=14)

        # Mission purpose
        team = 'Allemagne' if country.lower() == 'allemagne' else 'France'
        purpose = (
            f"Pour la mission : Formation professionnel, renforcer la coopération entre les deux équipes "
            f"(Tunis et {team}) et rencontre clients Planisware."
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, purpose, size=14)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature
        rep = config.company.representative_name or 'Mme ***REDACTED***'
        for line in [rep, 'Directrice Générale', 'PLW Tunisia']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._run(p, line, bold=True, size=14)

        buf = io.BytesIO()
        doc.save(buf)
        if output_path:
            doc.save(output_path)
        buf.seek(0)
        return buf
