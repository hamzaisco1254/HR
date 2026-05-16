"""
Lightweight RAG (Retrieval Augmented Generation) store — Postgres-backed.

Documents and chunks are persisted in rag_documents + rag_chunks tables.
Embeddings are stored as JSONB arrays (future work: migrate to pgvector
for in-DB cosine similarity search).

Until pgvector, we still compute cosine similarity in Python after
loading the full embedding set — which is fine for < ~5000 chunks.

Built-in reference documents (Code du travail, conventions collectives,
etc.) are pre-embedded at build time and shipped as a gzipped JSON
bundle (api/static_rag_bundle.json.gz). They are loaded once into
memory at module import and merged with DB chunks at search time, so
they are always available regardless of DB state.
"""
import gzip
import io
import json
import math
import os
import re
import uuid
from datetime import datetime
from typing import List, Optional

from gemini_client import embed_batch, embed_text
import db


# ─── Helpers ─────────────────────────────────────────────────────────

def cosine_similarity(a: list, b: list) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = na = nb = 0.0
    for x, y in zip(a, b):
        dot += x * y; na += x * x; nb += y * y
    if na == 0.0 or nb == 0.0:
        return 0.0
    return dot / (math.sqrt(na) * math.sqrt(nb))


def chunk_text(text: str, size: int = 900, overlap: int = 120) -> list:
    """Split text into overlapping, paragraph-aware chunks."""
    text = (text or '').strip()
    if not text:
        return []
    paras = [p.strip() for p in text.replace('\r', '').split('\n') if p.strip()]
    chunks, current = [], ''
    for para in paras:
        if len(current) + len(para) + 1 <= size:
            current = (current + '\n' + para).strip() if current else para
            continue
        if current:
            chunks.append(current)
        if len(para) > size:
            step = max(size - overlap, 200)
            for i in range(0, len(para), step):
                chunks.append(para[i:i + size])
            current = ''
        else:
            current = para
    if current:
        chunks.append(current)
    return chunks


def _clean_text(text: str) -> str:
    if not text:
        return ''
    text = text.replace('\x0c', '\n')
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Best-effort multilingual text extraction (AR / FR / EN)."""
    fn = (filename or '').lower()

    if fn.endswith('.pdf'):
        # PyMuPDF → pdfminer → pypdf
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype='pdf')
            parts = [p.get_text('text') for p in doc if (p.get_text('text') or '').strip()]
            doc.close()
            text = _clean_text('\n\n'.join(parts))
            if len(text) >= 50:
                return text
        except Exception:
            pass
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = _clean_text(pdfminer_extract(io.BytesIO(file_bytes)) or '')
            if len(text) >= 50:
                return text
        except Exception:
            pass
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            parts = []
            for page in reader.pages:
                try: txt = page.extract_text() or ''
                except: txt = ''
                if txt.strip(): parts.append(txt)
            text = _clean_text('\n\n'.join(parts))
            if len(text) >= 50:
                return text
        except Exception:
            pass
        return ''

    if fn.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            parts.append(cell.text.strip())
            return _clean_text('\n\n'.join(parts))
        except Exception:
            return ''

    if fn.endswith('.txt') or fn.endswith('.md'):
        try:
            return _clean_text(file_bytes.decode('utf-8', errors='ignore'))
        except Exception:
            return ''

    return ''


# ─── Store ───────────────────────────────────────────────────────────

CATEGORIES = {
    'convention':      'Convention collective',
    'law':             'Droit du travail',
    'policy':          'Politique interne',
    'contract':        'Contrat / Modele',
    'job_description': 'Fiche de poste',
    'other':           'Autre',
}


# ─── Static reference bundle ─────────────────────────────────────────
# Pre-embedded documents shipped with the codebase. Loaded once at
# import; merged with DB chunks in RagStore.search().

_STATIC_BUNDLE_PATH = os.path.join(os.path.dirname(__file__), 'static_rag_bundle.json.gz')

# Each item: {id, filename, title, category, size, chunk_count}
_STATIC_DOCS: list = []
# Each item: {doc_id, filename, category, idx, text, embedding}
_STATIC_CHUNKS: list = []


def _load_static_bundle() -> None:
    """Load the gzipped JSON reference bundle into memory.

    Silent no-op if the file is missing — lets local dev boot without
    a built bundle. Logs a warning if the file is malformed so build
    issues do not fail silently in production.
    """
    global _STATIC_DOCS, _STATIC_CHUNKS
    if not os.path.exists(_STATIC_BUNDLE_PATH):
        return
    try:
        with gzip.open(_STATIC_BUNDLE_PATH, 'rb') as f:
            bundle = json.loads(f.read().decode('utf-8'))
    except Exception as e:
        print(f'[rag_store] failed to load static bundle: {e}')
        return

    docs = bundle.get('documents') or []
    if not isinstance(docs, list):
        return

    for doc in docs:
        doc_id = doc.get('id')
        filename = doc.get('filename') or doc.get('title') or doc_id
        category = doc.get('category', 'other')
        if category not in CATEGORIES:
            category = 'other'
        chunks = doc.get('chunks') or []
        if not doc_id or not chunks:
            continue

        _STATIC_DOCS.append({
            'id':          doc_id,
            'filename':    filename,
            'title':       doc.get('title', filename),
            'category':    category,
            'size':        int(doc.get('size') or 0),
            'chunk_count': len(chunks),
        })
        for ch in chunks:
            emb = ch.get('embedding')
            text = ch.get('text') or ''
            if not isinstance(emb, list) or not text:
                continue
            _STATIC_CHUNKS.append({
                'doc_id':    doc_id,
                'filename':  filename,
                'category':  category,
                'idx':       int(ch.get('idx', 0)),
                'text':      text,
                'embedding': emb,
            })


_load_static_bundle()


class RagStore:
    def __init__(self, **_):
        pass  # schema bootstrapped by db_schema.py

    # -- Document management -----------------------------------------------

    def list_documents(self) -> list:
        # Built-in references first (always present, non-deletable)
        result = [
            {
                'id':             doc['id'],
                'filename':       doc['filename'],
                'title':          doc.get('title', doc['filename']),
                'category':       doc['category'],
                'category_label': CATEGORIES.get(doc['category'], 'Autre'),
                'size':           doc.get('size', 0),
                'chunks':         doc.get('chunk_count', 0),
                'created_at':     '',
                'is_builtin':     True,
            }
            for doc in _STATIC_DOCS
        ]

        rows = db.query("SELECT * FROM rag_documents ORDER BY created_at DESC")
        for doc in rows:
            chunk_cnt = (db.one(
                "SELECT COUNT(*) AS cnt FROM rag_chunks WHERE doc_id = %s", (doc['id'],)
            ) or {}).get('cnt', 0)
            result.append({
                'id':             doc['id'],
                'filename':       doc['filename'],
                'category':       doc.get('category', 'other'),
                'category_label': CATEGORIES.get(doc.get('category', 'other'), 'Autre'),
                'size':           doc.get('size', 0),
                'chunks':         chunk_cnt,
                'created_at':     str(doc.get('created_at') or ''),
                'is_builtin':     False,
            })
        return result

    def add_document(self, filename: str, file_bytes: bytes, category: str = 'other') -> dict:
        text = extract_text_from_file(file_bytes, filename)
        if not text or len(text.strip()) < 30:
            raise ValueError("Texte non extrait. Le fichier est vide, protege ou au format image non-OCR.")

        MAX_CHUNKS = 500
        chunk_size, overlap = 900, 120
        chunks = chunk_text(text, size=chunk_size, overlap=overlap)
        if len(chunks) > MAX_CHUNKS:
            ratio = len(chunks) / MAX_CHUNKS
            chunk_size = min(int(chunk_size * ratio), 4000)
            overlap = min(int(chunk_size * 0.12), 300)
            chunks = chunk_text(text, size=chunk_size, overlap=overlap)[:MAX_CHUNKS]

        if not chunks:
            raise ValueError("Aucun contenu indexable trouve.")

        try:
            all_embeddings = embed_batch(chunks, task_type='RETRIEVAL_DOCUMENT')
        except Exception as e:
            raise RuntimeError(f"Echec de l'indexation (embeddings) : {e}")

        if len(all_embeddings) != len(chunks):
            raise RuntimeError("Incoherence : nombre d'embeddings != nombre de chunks.")

        doc_id = uuid.uuid4().hex[:10]
        cat = category if category in CATEGORIES else 'other'

        db.execute(
            """INSERT INTO rag_documents (id, filename, category, size, created_at)
               VALUES (%s, %s, %s, %s, %s)""",
            (doc_id, filename, cat, len(file_bytes), datetime.utcnow()),
        )
        for i, (chunk, emb) in enumerate(zip(chunks, all_embeddings)):
            db.execute(
                """INSERT INTO rag_chunks (doc_id, idx, text, embedding)
                   VALUES (%s, %s, %s, %s::jsonb)""",
                (doc_id, i, chunk, json.dumps(emb)),
            )

        return {
            'id':       doc_id,
            'chunks':   len(chunks),
            'filename': filename,
            'category': cat,
        }

    def delete_document(self, doc_id: str) -> bool:
        # CASCADE deletes chunks automatically
        return db.execute("DELETE FROM rag_documents WHERE id = %s", (doc_id,)) > 0

    def stats(self) -> dict:
        doc_cnt   = (db.one("SELECT COUNT(*) AS cnt FROM rag_documents") or {}).get('cnt', 0)
        chunk_cnt = (db.one("SELECT COUNT(*) AS cnt FROM rag_chunks") or {}).get('cnt', 0)
        cat_rows  = db.query(
            "SELECT category, COUNT(*) AS cnt FROM rag_documents GROUP BY category"
        )
        by_cat = {r['category']: r['cnt'] for r in cat_rows}
        # Fold built-in references into the same totals
        for d in _STATIC_DOCS:
            by_cat[d['category']] = by_cat.get(d['category'], 0) + 1
        return {
            'total_documents': doc_cnt + len(_STATIC_DOCS),
            'total_chunks':    chunk_cnt + len(_STATIC_CHUNKS),
            'by_category':     by_cat,
            'builtin_documents': len(_STATIC_DOCS),
        }

    # -- Retrieval --------------------------------------------------------

    def search(self, query: str, top_k: int = 8, threshold: float = 0.0) -> list:
        if not query.strip():
            return []
        try:
            q_emb = embed_text(query, task_type='RETRIEVAL_QUERY')
        except Exception:
            return []

        # Load all chunks with their doc metadata
        # For performance at scale, this should become a pgvector <=> query
        rows = db.query("""
            SELECT c.doc_id, c.idx, c.text, c.embedding,
                   d.filename, d.category
            FROM rag_chunks c
            JOIN rag_documents d ON d.id = c.doc_id
        """)

        scored = []
        for r in rows:
            emb = r.get('embedding')
            if isinstance(emb, str):
                try: emb = json.loads(emb)
                except: continue
            if not isinstance(emb, list):
                continue
            sim = cosine_similarity(q_emb, emb)
            if sim >= threshold:
                scored.append({
                    'doc_id':         r['doc_id'],
                    'filename':       r['filename'],
                    'category':       r.get('category', 'other'),
                    'category_label': CATEGORIES.get(r.get('category', 'other'), 'Autre'),
                    'chunk_idx':      r.get('idx', 0),
                    'text':           r.get('text', ''),
                    'score':          sim,
                    'is_builtin':     False,
                })

        # Score built-in reference chunks against the same query embedding
        for c in _STATIC_CHUNKS:
            sim = cosine_similarity(q_emb, c['embedding'])
            if sim >= threshold:
                scored.append({
                    'doc_id':         c['doc_id'],
                    'filename':       c['filename'],
                    'category':       c['category'],
                    'category_label': CATEGORIES.get(c['category'], 'Autre'),
                    'chunk_idx':      c['idx'],
                    'text':           c['text'],
                    'score':          sim,
                    'is_builtin':     True,
                })

        scored.sort(key=lambda x: x['score'], reverse=True)
        return scored[:top_k]

    def has_documents(self) -> bool:
        if _STATIC_DOCS:
            return True
        cnt = (db.one("SELECT COUNT(*) AS cnt FROM rag_documents") or {}).get('cnt', 0)
        return cnt > 0
